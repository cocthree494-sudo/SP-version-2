"""Deterministic file text extraction and normalization."""

from __future__ import annotations

import io
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from pathlib import PurePosixPath
from typing import Any, Protocol

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.config import settings
from app.domains.knowledge.files import FileKind


class TextExtractionError(ValueError):
    """Safe deterministic parser failure persisted on source/job status."""

    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    title: str | None
    text: str
    metadata: dict[str, Any]


class TextExtractor(Protocol):
    def extract(self, data: bytes, *, filename: str) -> ExtractedDocument: ...


def normalize_text(value: str) -> str:
    """Normalize Unicode, line endings, controls, and excess blank lines."""

    normalized = unicodedata.normalize("NFC", value).replace("\r\n", "\n").replace("\r", "\n")
    normalized = "".join(
        character
        for character in normalized
        if character in {"\n", "\t"} or ord(character) >= 32
    )
    lines = [line.rstrip() for line in normalized.split("\n")]
    normalized = "\n".join(lines)
    normalized = re.sub(r"\n[ \t]*\n(?:[ \t]*\n)+", "\n\n", normalized)
    return normalized.strip()


def normalize_title(value: str | None) -> str | None:
    if value is None:
        return None
    normalized = " ".join(normalize_text(value).split())
    return normalized[:500] or None


def _require_useful_text(text: str) -> str:
    normalized = normalize_text(text)
    if not normalized:
        raise TextExtractionError("empty_document", "No useful text could be extracted")
    if len(normalized) > settings.FILE_PARSE_MAX_OUTPUT_CHARS:
        raise TextExtractionError("text_too_large", "Extracted text exceeds the processing limit")
    return normalized


class PdfTextExtractor:
    def extract(self, data: bytes, *, filename: str) -> ExtractedDocument:
        del filename
        try:
            reader = PdfReader(io.BytesIO(data), strict=False)
            if reader.is_encrypted:
                raise TextExtractionError("encrypted_pdf", "Encrypted PDFs are not supported")
            if len(reader.pages) > settings.FILE_PARSE_MAX_PAGES:
                raise TextExtractionError("pdf_page_limit", "PDF page count exceeds the limit")
            pages = [normalize_text(page.extract_text() or "") for page in reader.pages]
            metadata = reader.metadata
        except TextExtractionError:
            raise
        except Exception as exc:
            raise TextExtractionError("invalid_pdf", "The PDF could not be parsed") from exc
        text = _require_useful_text("\n\n".join(page for page in pages if page))
        title = normalize_title(str(metadata.title) if metadata and metadata.title else None)
        return ExtractedDocument(
            title=title,
            text=text,
            metadata={"page_count": len(reader.pages), "format": "pdf"},
        )


def _validate_docx_archive(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            names = {item.filename for item in infos}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise TextExtractionError("invalid_docx", "DOCX package files are missing")
            total_size = 0
            for item in infos:
                path = PurePosixPath(item.filename)
                if path.is_absolute() or ".." in path.parts:
                    raise TextExtractionError("unsafe_docx", "DOCX contains an unsafe path")
                total_size += item.file_size
                if total_size > settings.DOCX_MAX_UNCOMPRESSED_BYTES:
                    raise TextExtractionError(
                        "docx_too_large",
                        "DOCX expands beyond the safe limit",
                    )
                if item.compress_size and item.file_size / item.compress_size > 200:
                    raise TextExtractionError(
                        "unsafe_docx_compression",
                        "DOCX contains an unsafe compression ratio",
                    )
    except TextExtractionError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise TextExtractionError("invalid_docx", "The DOCX archive could not be parsed") from exc


class DocxTextExtractor:
    def extract(self, data: bytes, *, filename: str) -> ExtractedDocument:
        del filename
        _validate_docx_archive(data)
        try:
            document = DocxDocument(io.BytesIO(data))
            blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [normalize_text(cell.text) for cell in row.cells]
                    if any(cells):
                        blocks.append(" | ".join(cells))
            title = normalize_title(document.core_properties.title)
        except Exception as exc:
            raise TextExtractionError("invalid_docx", "The DOCX could not be parsed") from exc
        return ExtractedDocument(
            title=title,
            text=_require_useful_text("\n\n".join(blocks)),
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "format": "docx",
            },
        )


class PlainTextExtractor:
    def __init__(self, *, markdown: bool) -> None:
        self.markdown = markdown

    def extract(self, data: bytes, *, filename: str) -> ExtractedDocument:
        try:
            decoded = data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise TextExtractionError("invalid_text_encoding", "Text files must use UTF-8") from exc
        text = _require_useful_text(decoded)
        title = None
        if self.markdown:
            heading = re.search(r"(?m)^#{1,6}\s+(.+?)\s*$", text)
            title = normalize_title(heading.group(1) if heading else None)
        if title is None:
            title = normalize_title(PurePosixPath(filename).stem)
        return ExtractedDocument(
            title=title,
            text=text,
            metadata={"format": "markdown" if self.markdown else "text"},
        )


_EXTRACTORS: dict[FileKind, TextExtractor] = {
    FileKind.PDF: PdfTextExtractor(),
    FileKind.DOCX: DocxTextExtractor(),
    FileKind.TEXT: PlainTextExtractor(markdown=False),
    FileKind.MARKDOWN: PlainTextExtractor(markdown=True),
}


def extract_file(data: bytes, *, kind: FileKind, filename: str) -> ExtractedDocument:
    return _EXTRACTORS[kind].extract(data, filename=filename)


__all__ = [
    "ExtractedDocument",
    "TextExtractionError",
    "TextExtractor",
    "extract_file",
    "normalize_text",
    "normalize_title",
]
