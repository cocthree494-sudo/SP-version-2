"""Parser, chunking, embedding, and secure crawler fixtures."""

from __future__ import annotations

import io
import ipaddress
from pathlib import Path

import httpx
import pytest
from docx import Document as DocxDocument

from app.domains.knowledge.chunking import chunk_text
from app.domains.knowledge.crawler import CrawlError, WebsiteCrawler
from app.domains.knowledge.extraction import extract_file, normalize_text
from app.domains.knowledge.files import FileKind
from app.providers.embeddings import DeterministicEmbeddingProvider


class PublicResolver:
    async def resolve(
        self,
        _hostname: str,
    ) -> set[ipaddress.IPv4Address | ipaddress.IPv6Address]:
        return {ipaddress.ip_address("93.184.216.34")}


def minimal_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode())
        content.extend(obj)
        content.extend(b"\nendobj\n")
    xref = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode())
    content.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return bytes(content)


def docx_bytes() -> bytes:
    document = DocxDocument()
    document.core_properties.title = "Account Help"
    document.add_heading("Reset password", level=1)
    document.add_paragraph("Open settings and choose Reset password.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_all_parser_fixtures_extract_deterministic_text(tmp_path: Path) -> None:
    text_fixture = tmp_path / "guide.txt"
    text_fixture.write_text("First line\r\n\r\n\r\nSecond line  \r\n", encoding="utf-8")
    markdown = extract_file(
        b"# Refund policy\r\n\r\nRefunds take five days.\r\n",
        kind=FileKind.MARKDOWN,
        filename="refunds.md",
    )
    plain = extract_file(
        text_fixture.read_bytes(),
        kind=FileKind.TEXT,
        filename=text_fixture.name,
    )
    pdf = extract_file(minimal_pdf("Hello PDF"), kind=FileKind.PDF, filename="guide.pdf")
    docx = extract_file(docx_bytes(), kind=FileKind.DOCX, filename="guide.docx")

    assert markdown.title == "Refund policy"
    assert markdown.text == "# Refund policy\n\nRefunds take five days."
    assert plain.text == "First line\n\nSecond line"
    assert "Hello PDF" in pdf.text
    assert pdf.metadata["page_count"] == 1
    assert docx.title == "Account Help"
    assert "Reset password" in docx.text
    assert normalize_text("A\r\n\r\n\r\nB") == "A\n\nB"


@pytest.mark.asyncio
async def test_chunking_limits_overlap_and_deterministic_embeddings() -> None:
    text = "# Returns\n\n" + " ".join(f"word{index}" for index in range(80))
    chunks = chunk_text(text, max_tokens=20, overlap_tokens=5)
    provider = DeterministicEmbeddingProvider(dimensions=16)
    first = await provider.embed([chunk.content for chunk in chunks])
    second = await provider.embed([chunk.content for chunk in chunks])

    assert len(chunks) > 1
    assert all(chunk.token_count <= 20 for chunk in chunks)
    assert all(chunk.section == "Returns" for chunk in chunks)
    assert first.embeddings == second.embeddings
    assert all(len(embedding) == 16 for embedding in first.embeddings)


@pytest.mark.asyncio
async def test_crawler_honors_bounds_canonicalizes_and_deduplicates() -> None:
    async def handler(request: httpx.Request) -> httpx.Response:
        if request.url.path == "/robots.txt":
            return httpx.Response(200, text="User-agent: *\nAllow: /\n")
        if request.url.path == "/":
            return httpx.Response(
                200,
                headers={"content-type": "text/html"},
                text=(
                    "<html><head><title>Home</title></head><body>"
                    "<main><h1>Support</h1><p>Welcome.</p></main>"
                    '<a href="/help?utm_source=test">Help</a>'
                    '<a href="https://other.example/private">Outside</a>'
                    "</body></html>"
                ),
            )
        if request.url.path == "/help":
            return httpx.Response(
                200,
                headers={"content-type": "text/html"},
                text="<html><title>Help</title><main>Refund instructions.</main></html>",
            )
        return httpx.Response(404)

    async with httpx.AsyncClient(transport=httpx.MockTransport(handler)) as client:
        pages = await WebsiteCrawler(client, resolver=PublicResolver()).crawl(
            "https://example.com/",
            max_pages=2,
            max_depth=2,
            request_delay_seconds=0,
        )
    assert [page.url for page in pages] == [
        "https://example.com/",
        "https://example.com/help",
    ]
    assert pages[0].title == "Home"


@pytest.mark.asyncio
async def test_crawler_skips_rejected_child_page_without_hiding_start_failure() -> None:
    async def handler(request: httpx.Request) -> httpx.Response:
        if request.url.path == "/robots.txt":
            return httpx.Response(200, text="User-agent: *\nAllow: /\n")
        if request.url.path == "/":
            return httpx.Response(
                200,
                headers={"content-type": "text/html"},
                text=(
                    "<main>Useful home content.</main>"
                    '<a href="/missing">Broken link</a>'
                    '<a href="/help">Help</a>'
                ),
            )
        if request.url.path == "/help":
            return httpx.Response(
                200,
                headers={"content-type": "text/html"},
                text="<main>Useful help content.</main>",
            )
        return httpx.Response(404)

    async with httpx.AsyncClient(transport=httpx.MockTransport(handler)) as client:
        pages = await WebsiteCrawler(client, resolver=PublicResolver()).crawl(
            "https://example.com/",
            max_pages=3,
            max_depth=1,
            request_delay_seconds=0,
        )
    assert [page.url for page in pages] == [
        "https://example.com/",
        "https://example.com/help",
    ]

    async with httpx.AsyncClient(
        transport=httpx.MockTransport(lambda _request: httpx.Response(404))
    ) as client:
        with pytest.raises(CrawlError, match="HTTP 404"):
            await WebsiteCrawler(client, resolver=PublicResolver()).crawl(
                "https://example.com/",
                max_pages=1,
                max_depth=0,
                request_delay_seconds=0,
            )


@pytest.mark.asyncio
async def test_crawler_blocks_redirect_to_private_network() -> None:
    async def handler(request: httpx.Request) -> httpx.Response:
        if request.url.path == "/robots.txt":
            return httpx.Response(200, text="User-agent: *\nAllow: /\n")
        return httpx.Response(302, headers={"location": "http://127.0.0.1/private"})

    async with httpx.AsyncClient(transport=httpx.MockTransport(handler)) as client:
        with pytest.raises(CrawlError) as captured:
            await WebsiteCrawler(client, resolver=PublicResolver()).crawl(
                "https://example.com/",
                max_pages=1,
                max_depth=0,
                request_delay_seconds=0,
            )
    assert captured.value.code == "unsafe_host"
    assert captured.value.retryable is False
